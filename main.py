import docx, fitz, os
from PIL import Image
from docx2pdf import convert

# # create new document
# doc = docx.Document()

# # write each character on the document
# print("writing the document...")
# p = doc.add_paragraph()
# run = p.add_run("ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789 ")
# run.font.name = "Candara"
# run.font.size = docx.shared.Pt(520)

# # save the document
# doc.save("out.docx")
# # and convert it to pdf
# print("converting it to pdf...")
# convert("out.docx", "out.pdf")

# open that pdf file
doc = fitz.open("out.pdf")

# and convert it to pngs
print("converting to png...")
for i, page in enumerate(doc):
    pix = page.get_pixmap()

    pix.save(f"imgs/{i}.png")


# crop the images and get the size of the sprite
spriteSize = [0, 0]
leastHeight = float('INF')
leastWidth = float('INF')
lastI = 0
for i in range(len(doc)):
    img = Image.open(f"imgs/{i}.png")
    img = img.convert("RGBA")

    matrix = img.getdata()
    newMatrix = []

    # remove the white pixels from each of the png
    for pixel in matrix: # pixel --> r b b a
        # removing the white pixles
        if pixel[0] == 255 and pixel[1] == 255 and pixel[2] == 255:
            newMatrix.append((255,255,255,0))
        # changing the black pixels into white
        elif pixel[0] == 0 and pixel[1] == 0 and pixel[2] == 0 and pixel[3] == 255:
            newMatrix.append((255,255,255,255))
        else:
            newMatrix.append((0,0,0,0))

    img.putdata(newMatrix)

    # get the rect bounds of each image
    leftmost = img.width
    rightmost = 0
    topmost = img.height
    botmost = 0

    for r in range(img.height):
        for c in range(img.width):
            pixel = img.getpixel( (c, r) )
            
            if pixel[3] == 255 and pixel[0] == 255 and pixel[1] == 255 and pixel[2] == 255:
                if c < leftmost:
                    leftmost = c
                elif c > rightmost:
                    rightmost = c

                if r < topmost:
                    topmost = r
                elif r > botmost:
                    botmost = r

    # crop the image based on the bounds
    img = img.crop((leftmost, topmost, rightmost, botmost)) # box is (left, top, right, bot)

    # save the width and height of the sprite
    spriteSize[0] += (rightmost - leftmost) + 1 # +1 for the letter separator
    if (botmost - topmost) > spriteSize[1]:
        spriteSize[1] = (botmost - topmost)

    # save the least height amongst the upper case letters
    if i < 25: # these are the capital case letters
        if (botmost - topmost) < leastHeight:
            leastHeight = (botmost - topmost)
    
    # save the index in the lastI varibale
    lastI = i
    
    # save the least width amongst the letters
    if (rightmost - leftmost) < leastWidth:
        leastWidth = (rightmost - leftmost)

    # save it
    img.save(f"final imgs/{i}.png", "PNG")

    os.system("cls")
    print(f"cropping each letter... {int(i/len(doc)*100)}% complete")

# add an empty image to signify space
space = Image.new("RGBA", (leastWidth, spriteSize[1]), (0,0,0,0))
space.save(f"final imgs/{lastI+1}.png")
spriteSize[0] += space.width + 1 # +1 for the letter separator

# build the sprite
alphabet = "ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789 "
sinkingLowerCases = "gjpqy"
sprite = Image.new("RGBA", spriteSize, (0,0,0,0))
separator = Image.new("RGBA", (1, sprite.height), (118, 118, 118, 255))

left = 0
top = 0
bot = 0
right = 0
lowerCaseTop = 0

for i in range(len(doc) + 1): # +1 to compensate for the space
    img = Image.open(f"final imgs/{i}.png")

    right = left + img.width
    if i < 25: # i < 25 means the upper cases
        top = 0
        bot = top + img.height
    else:
        bot = leastHeight
        top = bot - img.height
    
    if alphabet[i] in sinkingLowerCases:
        bot = sprite.height
        top = bot - img.height

    # place the letter image on the sprite
    sprite.paste(img, (left, top, right, bot))
    # update the left
    left += img.width + 1

    os.system("cls")
    print(f"building the sprite... {int(i/len(doc)*100)}% complete")


# place the separators
left = 0
right = 0

for i in range(len(doc) + 1): #+1 to compensate for the space
    img = Image.open(f"final imgs/{i}.png")

    right = left + img.width

    # place the separator on the sprite
    sprite.paste(separator, (right, 0, right+1, sprite.height))
    # update the left
    left += img.width + 1


# save the sprite
sprite.save("sprite.png")
print("done")
os.system("pause > nul")