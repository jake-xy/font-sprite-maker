import docx, os
from docx2pdf import convert

# create new document
doc = docx.Document()

fontName = input("font name > ")
fontSize = int(input("font size (pt) > "))

# write each character on the document
print("writing the document...")
p = doc.add_paragraph()
run = p.add_run("ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789 ")
run.font.name = fontName
run.font.size = docx.shared.Pt(fontSize)

# save the document
doc.save("out.docx")

print("done")
os.system("pause > nul")
